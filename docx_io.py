from __future__ import annotations

import base64
import io
import re
import zipfile
import xml.etree.ElementTree as ET

from docx import Document as WordDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as docx_qn
from docx.shared import Inches, Pt


XML_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
}

for prefix, uri in XML_NS.items():
    ET.register_namespace(prefix, uri)


NUMBER_FORMATS = {"decimal", "upperLetter", "lowerLetter", "upperRoman", "lowerRoman", "bullet"}
MAX_NUMBERING_LEVEL = 8
DEFAULT_LATIN_FONT = "Times New Roman"
DEFAULT_EAST_ASIA_FONT = "SimSun"
DEFAULT_FONT_FAMILY = f"{DEFAULT_LATIN_FONT}, {DEFAULT_EAST_ASIA_FONT}"
DEFAULT_LINE_SPACING = 1.5
ALLOWED_STYLE_ORDER = ("Normal", "Heading1", "Heading2", "Heading3")
ALLOWED_STYLE_IDS = set(ALLOWED_STYLE_ORDER)
STYLE_ALIAS_MAP = {
    "normal": "Normal",
    "正文": "Normal",
    "h1": "Heading1",
    "heading1": "Heading1",
    "heading 1": "Heading1",
    "标题1": "Heading1",
    "标题 1": "Heading1",
    "h2": "Heading2",
    "heading2": "Heading2",
    "heading 2": "Heading2",
    "标题2": "Heading2",
    "标题 2": "Heading2",
    "h3": "Heading3",
    "heading3": "Heading3",
    "heading 3": "Heading3",
    "标题3": "Heading3",
    "标题 3": "Heading3",
}


def qn(prefix: str, name: str) -> str:
    return f"{{{XML_NS[prefix]}}}{name}"


def _normalize_num_fmt(value: object) -> str:
    raw = str(value or "decimal").strip()
    return raw if raw in NUMBER_FORMATS else "decimal"


def _normalize_ilvl(value: object) -> int:
    try:
        return max(min(int(str(value)), MAX_NUMBERING_LEVEL), 0)
    except (TypeError, ValueError):
        return 0


def _default_lvl_text(ilvl: int) -> str:
    ilvl = _normalize_ilvl(ilvl)
    if ilvl == 0:
        return "%1."
    return ".".join(f"%{index + 1}" for index in range(ilvl + 1)) + "."


def _normalize_numbering(payload: dict | None) -> dict | None:
    if not isinstance(payload, dict):
        return None
    if payload.get("enabled") is False:
        return None
    ilvl = _normalize_ilvl(payload.get("ilvl", payload.get("level", 0)))
    try:
        start = max(int(str(payload.get("start", 1))), 1)
    except (TypeError, ValueError):
        start = 1
    num_fmt = _normalize_num_fmt(payload.get("num_fmt"))
    default_text = "•" if num_fmt == "bullet" else _default_lvl_text(ilvl)
    lvl_text = str(payload.get("lvl_text") or payload.get("level_text") or "").strip() or default_text
    return {
        "enabled": True,
        "list_id": str(payload.get("list_id") or payload.get("listId") or payload.get("num_id") or "").strip(),
        "ilvl": ilvl,
        "num_fmt": num_fmt,
        "lvl_text": lvl_text,
        "start": start,
    }


def _normalize_style_numbering(payload: dict | None) -> dict | None:
    if not isinstance(payload, dict):
        return None
    num_id = str(payload.get("num_id") or payload.get("numId") or "").strip()
    if not num_id:
        return None
    return {"num_id": num_id, "ilvl": _normalize_ilvl(payload.get("ilvl", 0))}


def _iter_paragraph_blocks(blocks: list[dict]):
    for block in blocks:
        if not isinstance(block, dict):
            continue
        block_type = block.get("type")
        if block_type == "paragraph":
            yield block
            continue
        if block_type != "table":
            continue
        for row in block.get("rows") or []:
            for cell in row or []:
                if isinstance(cell, dict):
                    yield from _iter_paragraph_blocks(cell.get("paragraphs") or [])


def _collect_numbering_lists(blocks: list[dict]) -> tuple[list[str], dict[str, dict[int, dict[str, object]]]]:
    list_levels: dict[str, dict[int, dict[str, object]]] = {}
    ordered_list_ids: list[str] = []
    auto_index = 1

    for paragraph in _iter_paragraph_blocks(blocks):
        numbering = _normalize_numbering(paragraph.get("numbering"))
        if numbering is None:
            continue
        list_id = numbering["list_id"] or f"list-{auto_index}"
        if not numbering["list_id"]:
            auto_index += 1
            numbering["list_id"] = list_id
            paragraph["numbering"] = numbering
        if list_id not in list_levels:
            list_levels[list_id] = {}
            ordered_list_ids.append(list_id)
        ilvl = numbering["ilvl"]
        list_levels[list_id].setdefault(
            ilvl,
            {"start": numbering["start"], "num_fmt": numbering["num_fmt"], "lvl_text": numbering["lvl_text"]},
        )
    return ordered_list_ids, list_levels


def _parse_num_suffix(list_id: str) -> str | None:
    if list_id.startswith("num-"):
        suffix = list_id[4:].strip()
        if suffix.isdigit():
            return suffix
    return None


def _max_attr_int(nodes: list[ET.Element], attr_qname: str) -> int:
    values: list[int] = []
    for node in nodes:
        raw = node.attrib.get(attr_qname)
        if raw and str(raw).isdigit():
            values.append(int(raw))
    return max(values) if values else 0


def _append_numbering_definition(
    root: ET.Element,
    num_id: str,
    abstract_num_id: str,
    levels: dict[int, dict[str, object]],
) -> None:
    max_level = max(levels.keys()) if levels else 0
    abstract_num = ET.SubElement(root, qn("w", "abstractNum"), {qn("w", "abstractNumId"): abstract_num_id})
    ET.SubElement(abstract_num, qn("w", "multiLevelType"), {qn("w", "val"): "hybridMultilevel"})

    for ilvl in range(max_level + 1):
        config = levels.get(ilvl, {"start": 1, "num_fmt": "decimal", "lvl_text": _default_lvl_text(ilvl)})
        num_fmt = _normalize_num_fmt(config.get("num_fmt"))
        try:
            start = max(int(str(config.get("start", 1))), 1)
        except (TypeError, ValueError):
            start = 1
        lvl = ET.SubElement(abstract_num, qn("w", "lvl"), {qn("w", "ilvl"): str(ilvl)})
        ET.SubElement(lvl, qn("w", "start"), {qn("w", "val"): str(start)})
        ET.SubElement(lvl, qn("w", "numFmt"), {qn("w", "val"): num_fmt})
        default_text = "•" if num_fmt == "bullet" else _default_lvl_text(ilvl)
        ET.SubElement(lvl, qn("w", "lvlText"), {qn("w", "val"): str(config.get("lvl_text") or default_text)})
        if num_fmt == "bullet":
            ET.SubElement(lvl, qn("w", "rPr"))
            lvl_rpr = lvl.find(qn("w", "rPr"))
            ET.SubElement(lvl_rpr, qn("w", "rFonts"), {qn("w", "ascii"): "Symbol", qn("w", "hAnsi"): "Symbol"})
        ET.SubElement(lvl, qn("w", "lvlJc"), {qn("w", "val"): "left"})
        p_pr = ET.SubElement(lvl, qn("w", "pPr"))
        ET.SubElement(
            p_pr,
            qn("w", "ind"),
            {qn("w", "left"): str((ilvl + 1) * 720), qn("w", "hanging"): "360"},
        )

    num_node = ET.SubElement(root, qn("w", "num"), {qn("w", "numId"): num_id})
    ET.SubElement(num_node, qn("w", "abstractNumId"), {qn("w", "val"): abstract_num_id})


def _build_numbering_payload(
    blocks: list[dict],
    preserved_numbering_xml: bytes | None = None,
    preserve_existing: bool = False,
) -> tuple[dict[str, str], bytes | None]:
    ordered_list_ids, list_levels = _collect_numbering_lists(blocks)

    if preserved_numbering_xml and preserve_existing:
        parsed_from_preserved = True
        try:
            root = ET.fromstring(preserved_numbering_xml)
        except ET.ParseError:
            parsed_from_preserved = False
            root = ET.Element(qn("w", "numbering"))
        list_id_to_num_id: dict[str, str] = {}
        existing_num_ids: set[str] = set()
        for num_node in root.findall(qn("w", "num")):
            num_id = str(num_node.attrib.get(qn("w", "numId")) or "").strip()
            if not num_id:
                continue
            existing_num_ids.add(num_id)
            list_id_to_num_id[f"num-{num_id}"] = num_id

        next_num_id = _max_attr_int(root.findall(qn("w", "num")), qn("w", "numId"))
        next_abstract_id = _max_attr_int(root.findall(qn("w", "abstractNum")), qn("w", "abstractNumId"))
        appended = False

        for list_id in ordered_list_ids:
            existing_num = _parse_num_suffix(list_id)
            if existing_num and existing_num in existing_num_ids:
                list_id_to_num_id[list_id] = existing_num
                continue
            next_num_id += 1
            next_abstract_id += 1
            num_id = str(existing_num or next_num_id)
            while num_id in existing_num_ids:
                next_num_id += 1
                num_id = str(next_num_id)
            existing_num_ids.add(num_id)
            levels = list_levels.get(list_id, {})
            _append_numbering_definition(root, num_id, str(next_abstract_id), levels)
            appended = True
            list_id_to_num_id[list_id] = num_id

        if parsed_from_preserved and not appended:
            return list_id_to_num_id, preserved_numbering_xml
        return list_id_to_num_id, ET.tostring(root, encoding="utf-8", xml_declaration=True)

    if not ordered_list_ids:
        if preserved_numbering_xml and preserve_existing:
            return {}, preserved_numbering_xml
        return {}, None

    root = ET.Element(qn("w", "numbering"))
    list_id_to_num_id: dict[str, str] = {}
    for index, list_id in enumerate(ordered_list_ids, start=1):
        num_id = str(index)
        abstract_num_id = str(index)
        _append_numbering_definition(root, num_id, abstract_num_id, list_levels[list_id])
        list_id_to_num_id[list_id] = num_id

    return list_id_to_num_id, ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _parse_level_node(level_node: ET.Element, fallback_ilvl: int = 0) -> tuple[int, dict]:
    ilvl = _normalize_ilvl(level_node.attrib.get(qn("w", "ilvl"), fallback_ilvl))
    start = 1
    start_node = level_node.find(qn("w", "start"))
    if start_node is not None:
        try:
            start = max(int(start_node.attrib.get(qn("w", "val"), "1")), 1)
        except ValueError:
            start = 1
    num_fmt_node = level_node.find(qn("w", "numFmt"))
    num_fmt = _normalize_num_fmt(num_fmt_node.attrib.get(qn("w", "val")) if num_fmt_node is not None else "decimal")
    lvl_text_node = level_node.find(qn("w", "lvlText"))
    lvl_text = (lvl_text_node.attrib.get(qn("w", "val")) if lvl_text_node is not None else "") or _default_lvl_text(ilvl)
    if num_fmt == "bullet" and lvl_text in {"", "?"}:
        lvl_text = "•"
    return ilvl, {"start": start, "num_fmt": num_fmt, "lvl_text": lvl_text}


def _parse_numbering(archive: zipfile.ZipFile) -> dict[str, dict[int, dict]]:
    try:
        root = ET.fromstring(archive.read("word/numbering.xml"))
    except (KeyError, ET.ParseError):
        return {}

    abstract_levels: dict[str, dict[int, dict]] = {}
    for abstract_node in root.findall(qn("w", "abstractNum")):
        abstract_id = abstract_node.attrib.get(qn("w", "abstractNumId"))
        if not abstract_id:
            continue
        levels: dict[int, dict] = {}
        for level_node in abstract_node.findall(qn("w", "lvl")):
            ilvl, config = _parse_level_node(level_node)
            levels[ilvl] = config
        abstract_levels[abstract_id] = levels

    numbering: dict[str, dict[int, dict]] = {}
    for num_node in root.findall(qn("w", "num")):
        num_id = num_node.attrib.get(qn("w", "numId"))
        if not num_id:
            continue
        abstract_ref = num_node.find(qn("w", "abstractNumId"))
        abstract_id = abstract_ref.attrib.get(qn("w", "val")) if abstract_ref is not None else None
        levels = {ilvl: dict(config) for ilvl, config in (abstract_levels.get(abstract_id or "") or {}).items()}

        for override_node in num_node.findall(qn("w", "lvlOverride")):
            ilvl = _normalize_ilvl(override_node.attrib.get(qn("w", "ilvl"), 0))
            if ilvl not in levels:
                levels[ilvl] = {"start": 1, "num_fmt": "decimal", "lvl_text": _default_lvl_text(ilvl)}
            start_override = override_node.find(qn("w", "startOverride"))
            if start_override is not None:
                try:
                    levels[ilvl]["start"] = max(int(start_override.attrib.get(qn("w", "val"), "1")), 1)
                except ValueError:
                    levels[ilvl]["start"] = 1
            level_override = override_node.find(qn("w", "lvl"))
            if level_override is not None:
                _, config = _parse_level_node(level_override, ilvl)
                levels[ilvl] = config

        numbering[num_id] = levels
    return numbering


def _default_descriptor() -> list:
    return [DEFAULT_FONT_FAMILY, 12, False, False, False, ""]


def _canonical_style_id(value: object) -> str | None:
    raw = str(value or "").strip()
    if not raw:
        return None
    compact = re.sub(r"\s+", "", raw).lower()
    lowered = raw.lower()
    return STYLE_ALIAS_MAP.get(lowered) or STYLE_ALIAS_MAP.get(compact)


def _normalize_descriptor(descriptor: list | tuple | None) -> list:
    values = list((descriptor or _default_descriptor())[:6])
    values += _default_descriptor()[len(values):]
    family, size, bold, italic, underline, bg_color = values[:6]
    family = str(family or DEFAULT_FONT_FAMILY)
    size = max(int(round(float(size or 12))), 1)
    bg_color = _normalize_color(bg_color)
    return [family, size, bool(bold), bool(italic), bool(underline), bg_color]


def _normalize_color(value: object) -> str:
    if value is None:
        return ""
    raw = str(value).strip().lower()
    if not raw or raw == "transparent" or raw == "none":
        return ""
    if raw.startswith("#"):
        if len(raw) == 4:
            return f"#{raw[1]}{raw[1]}{raw[2]}{raw[2]}{raw[3]}{raw[3]}"
        return raw if len(raw) == 7 else ""
    rgb_match = re.match(r"rgba?\(([^)]+)\)", raw)
    if rgb_match:
        parts = [item.strip() for item in rgb_match.group(1).split(",")]
        try:
            r = int(float(parts[0]))
            g = int(float(parts[1]))
            b = int(float(parts[2]))
            a = float(parts[3]) if len(parts) > 3 else 1.0
        except (ValueError, IndexError):
            return ""
        if a == 0:
            return ""
        r = max(0, min(255, r))
        g = max(0, min(255, g))
        b = max(0, min(255, b))
        return f"#{r:02x}{g:02x}{b:02x}"
    return ""


def _split_font_family(family: object) -> tuple[str, str]:
    parts = [part.strip().strip("\"'") for part in str(family or DEFAULT_FONT_FAMILY).split(",") if part.strip()]
    latin = parts[0] if parts else DEFAULT_LATIN_FONT
    east_asia = parts[1] if len(parts) > 1 else (DEFAULT_EAST_ASIA_FONT if latin == DEFAULT_LATIN_FONT else latin)
    return latin, east_asia


def _rfonts_attrs(family: object) -> dict[str, str]:
    latin, east_asia = _split_font_family(family)
    return {
        qn("w", "ascii"): latin,
        qn("w", "hAnsi"): latin,
        qn("w", "eastAsia"): east_asia,
    }


def _set_docx_rfonts(r_pr, family: object) -> None:
    latin, east_asia = _split_font_family(family)
    r_fonts = r_pr.find(docx_qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(docx_qn("w:ascii"), latin)
    r_fonts.set(docx_qn("w:hAnsi"), latin)
    r_fonts.set(docx_qn("w:eastAsia"), east_asia)


def _toggle_property_enabled(node: ET.Element | None, *, none_is_false: bool = False) -> bool:
    if node is None:
        return False
    raw = node.attrib.get(qn("w", "val"))
    if raw is None:
        return True
    normalized = str(raw).strip().lower()
    false_values = {"0", "false", "off"}
    if none_is_false:
        false_values.add("none")
    return normalized not in false_values


_HIGHLIGHT_MAP = {
    "yellow": "#ffff00",
    "green": "#00ff00",
    "cyan": "#00ffff",
    "magenta": "#ff00ff",
    "blue": "#0000ff",
    "red": "#ff0000",
    "darkBlue": "#00008b",
    "darkCyan": "#008b8b",
    "darkGreen": "#006400",
    "darkMagenta": "#8b008b",
    "darkRed": "#8b0000",
    "darkYellow": "#b8860b",
    "darkGray": "#a9a9a9",
    "lightGray": "#d3d3d3",
    "black": "#000000",
    "white": "#ffffff",
}


def _highlight_to_hex(value: object) -> str:
    if value is None:
        return ""
    raw = str(value)
    if not raw:
        return ""
    return _HIGHLIGHT_MAP.get(raw, "")


def _make_style(
    style_id: str,
    name: str,
    descriptor: list | None = None,
    alignment: str = "left",
    outline_level: int | None = None,
    is_default: bool = False,
    based_on: str | None = None,
    line_spacing: float = DEFAULT_LINE_SPACING,
    space_before: int = 0,
    space_after: int = 0,
    numbering: dict | None = None,
) -> dict:
    return {
        "id": style_id,
        "name": name,
        "type": "paragraph",
        "descriptor": _normalize_descriptor(descriptor),
        "alignment": alignment if alignment in {"left", "center", "right", "justify"} else "left",
        "outline_level": outline_level if outline_level in {0, 1, 2} else None,
        "is_default": bool(is_default),
        "based_on": based_on or None,
        "line_spacing": _normalize_line_spacing(line_spacing, DEFAULT_LINE_SPACING),
        "space_before": _normalize_spacing(space_before, 0),
        "space_after": _normalize_spacing(space_after, 0),
        "numbering": _normalize_style_numbering(numbering),
    }


def _normalize_spacing(value: object, fallback: int) -> int:
    try:
        raw = fallback if value is None else value
        return max(int(round(float(str(raw)))), 0)
    except (TypeError, ValueError):
        return fallback


def _normalize_line_spacing(value: object, fallback: float = DEFAULT_LINE_SPACING) -> float:
    try:
        raw = fallback if value is None else value
        return max(float(str(raw)), 1.0)
    except (TypeError, ValueError):
        return fallback


def _builtin_styles() -> list[dict]:
    return [
        _make_style("Normal", "Normal", [DEFAULT_FONT_FAMILY, 12, False, False, False], is_default=True),
        _make_style("Heading1", "Heading 1", [DEFAULT_FONT_FAMILY, 20, True, False, False], outline_level=0, based_on="Normal"),
        _make_style("Heading2", "Heading 2", [DEFAULT_FONT_FAMILY, 16, True, False, False], outline_level=1, based_on="Normal"),
        _make_style("Heading3", "Heading 3", [DEFAULT_FONT_FAMILY, 14, True, False, False], outline_level=2, based_on="Normal"),
    ]


def _style_map(styles: list[dict]) -> dict[str, dict]:
    return {style["id"]: style for style in styles if style.get("id")}


def _normalize_styles(payload: dict | None) -> list[dict]:
    items = ((payload or {}).get("paragraph") or []) if isinstance(payload, dict) else []
    incoming: dict[str, dict] = {}
    for style in items:
        if not isinstance(style, dict):
            continue
        style_id = _canonical_style_id(style.get("id"))
        if style_id in ALLOWED_STYLE_IDS:
            incoming[style_id] = style

    merged: list[dict] = []
    for base_style in _builtin_styles():
        style_id = base_style["id"]
        style = incoming.get(style_id) or {}
        merged.append(
            _make_style(
                style_id,
                base_style["name"],
                style.get("descriptor", base_style.get("descriptor")),
                str(style.get("alignment") or base_style.get("alignment") or "left"),
                base_style.get("outline_level"),
                bool(base_style.get("is_default")),
                base_style.get("based_on"),
                style.get("line_spacing", base_style.get("line_spacing", DEFAULT_LINE_SPACING)),
                style.get("space_before", base_style.get("space_before", 0)),
                style.get("space_after", base_style.get("space_after", 0)),
                style.get("numbering"),
            )
        )
    return merged


def _safe_style_id(name: str, used: set[str]) -> str:
    base = re.sub(r"[^A-Za-z0-9]+", "", name.title()) or "CustomStyle"
    candidate = base
    index = 2
    while candidate in used:
        candidate = f"{base}{index}"
        index += 1
    return candidate


def _clear_imported_paragraph_format(block: dict) -> dict:
    block["style"] = "normal"
    block["style_id"] = "Normal"
    block["style_name"] = "Normal"
    block["alignment"] = "align_left"
    block["line_spacing"] = DEFAULT_LINE_SPACING
    block["space_before"] = 0
    block["space_after"] = 0
    block.pop("numbering", None)
    default_descriptor = _default_descriptor()
    for run in block.get("runs") or []:
        run["descriptor"] = default_descriptor
    return block


def _style_ref_for_block(block: dict, styles: dict[str, dict]) -> str:
    style_id = _canonical_style_id(block.get("style_id"))
    if style_id and style_id in styles:
        return style_id
    style_key = str(block.get("style") or "")
    return {
        "heading1": "Heading1",
        "heading2": "Heading2",
        "heading3": "Heading3",
    }.get(style_key, "Normal")


def _style_to_block_style(style: dict | None) -> str:
    if not style:
        return "normal"
    outline_level = style.get("outline_level")
    if outline_level == 0:
        return "heading1"
    if outline_level == 1:
        return "heading2"
    if outline_level == 2:
        return "heading3"
    return "normal"


def _default_styles_payload() -> dict:
    return {"paragraph": _builtin_styles()}


def _root_rels_xml() -> bytes:
    root = ET.Element(qn("pr", "Relationships"))
    ET.SubElement(
        root,
        qn("pr", "Relationship"),
        {
            "Id": "rId1",
            "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument",
            "Target": "word/document.xml",
        },
    )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _content_types_xml(image_exts: set[str], has_numbering: bool = False) -> bytes:
    root = ET.Element("Types", xmlns="http://schemas.openxmlformats.org/package/2006/content-types")
    ET.SubElement(root, "Default", Extension="rels", ContentType="application/vnd.openxmlformats-package.relationships+xml")
    ET.SubElement(root, "Default", Extension="xml", ContentType="application/xml")
    mime_map = {
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "gif": "image/gif",
        "bmp": "image/bmp",
        "webp": "image/webp",
    }
    for ext in sorted(image_exts):
        ET.SubElement(root, "Default", Extension=ext, ContentType=mime_map.get(ext, "image/png"))
    ET.SubElement(
        root,
        "Override",
        PartName="/word/document.xml",
        ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
    )
    ET.SubElement(
        root,
        "Override",
        PartName="/word/styles.xml",
        ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml",
    )
    if has_numbering:
        ET.SubElement(
            root,
            "Override",
            PartName="/word/numbering.xml",
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
        )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _append_run_properties(parent: ET.Element, descriptor: list) -> None:
    family, size, bold, italic, underline, bg_color = _normalize_descriptor(descriptor)
    ET.SubElement(
        parent,
        qn("w", "rFonts"),
        _rfonts_attrs(family),
    )
    ET.SubElement(parent, qn("w", "sz"), {qn("w", "val"): str(size * 2)})
    if bold:
        ET.SubElement(parent, qn("w", "b"))
    if italic:
        ET.SubElement(parent, qn("w", "i"))
    if underline:
        ET.SubElement(parent, qn("w", "u"), {qn("w", "val"): "single"})
    if bg_color:
        ET.SubElement(
            parent,
            qn("w", "shd"),
            {qn("w", "val"): "clear", qn("w", "color"): "auto", qn("w", "fill"): bg_color.lstrip("#")},
        )


def _styles_xml(styles_payload: dict | None = None) -> bytes:
    styles = _normalize_styles(styles_payload)
    root = ET.Element(qn("w", "styles"))
    for style in styles:
        attrs = {qn("w", "type"): "paragraph", qn("w", "styleId"): style["id"]}
        if style.get("is_default"):
            attrs[qn("w", "default")] = "1"
        style_node = ET.SubElement(root, qn("w", "style"), attrs)
        ET.SubElement(style_node, qn("w", "name"), {qn("w", "val"): style["name"]})
        if style.get("based_on"):
            ET.SubElement(style_node, qn("w", "basedOn"), {qn("w", "val"): style["based_on"]})
        p_pr = ET.SubElement(style_node, qn("w", "pPr"))
        alignment = style.get("alignment") or "left"
        ET.SubElement(p_pr, qn("w", "jc"), {qn("w", "val"): "both" if alignment == "justify" else alignment})
        outline_level = style.get("outline_level")
        if outline_level is not None:
            ET.SubElement(p_pr, qn("w", "outlineLvl"), {qn("w", "val"): str(outline_level)})
        style_numbering = _normalize_style_numbering(style.get("numbering"))
        if style_numbering is not None:
            num_pr = ET.SubElement(p_pr, qn("w", "numPr"))
            ET.SubElement(num_pr, qn("w", "ilvl"), {qn("w", "val"): str(style_numbering["ilvl"])})
            ET.SubElement(num_pr, qn("w", "numId"), {qn("w", "val"): style_numbering["num_id"]})
        ET.SubElement(
            p_pr,
            qn("w", "spacing"),
            {
                qn("w", "before"): str(_normalize_spacing(style.get("space_before"), 0) * 20),
                qn("w", "after"): str(_normalize_spacing(style.get("space_after"), 0) * 20),
                qn("w", "line"): str(int(round(_normalize_line_spacing(style.get("line_spacing"), DEFAULT_LINE_SPACING) * 240))),
                qn("w", "lineRule"): "auto",
            },
        )
        r_pr = ET.SubElement(style_node, qn("w", "rPr"))
        _append_run_properties(r_pr, style.get("descriptor") or _default_descriptor())
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _document_rels_xml(entries: list[tuple[str, str]], has_numbering: bool = False) -> bytes:
    root = ET.Element(qn("pr", "Relationships"))
    ET.SubElement(
        root,
        qn("pr", "Relationship"),
        {
            "Id": "rIdStyles",
            "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles",
            "Target": "styles.xml",
        },
    )
    if has_numbering:
        ET.SubElement(
            root,
            qn("pr", "Relationship"),
            {
                "Id": "rIdNumbering",
                "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering",
                "Target": "numbering.xml",
            },
        )
    for rel_id, target in entries:
        ET.SubElement(
            root,
            qn("pr", "Relationship"),
            {
                "Id": rel_id,
                "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
                "Target": target,
            },
        )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _run_node(text: str, descriptor: list) -> ET.Element:
    run = ET.Element(qn("w", "r"))
    r_pr = ET.SubElement(run, qn("w", "rPr"))
    _append_run_properties(r_pr, descriptor)
    text_node = ET.SubElement(run, qn("w", "t"))
    if text.startswith(" ") or text.endswith(" "):
        text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = text
    return run


def _paragraph_node(block: dict, styles: dict[str, dict], list_id_to_num_id: dict[str, str]) -> ET.Element:
    paragraph = ET.Element(qn("w", "p"))
    p_pr = ET.SubElement(paragraph, qn("w", "pPr"))
    style_id = _style_ref_for_block(block, styles)
    ET.SubElement(p_pr, qn("w", "pStyle"), {qn("w", "val"): style_id})
    block_alignment = str(block.get("alignment") or "")
    if block_alignment == "align_center":
        alignment = "center"
    elif block_alignment == "align_right":
        alignment = "right"
    elif block_alignment == "align_justify":
        alignment = "justify"
    elif block_alignment == "align_left":
        alignment = "left"
    else:
        alignment = str(styles.get(style_id, {}).get("alignment") or "left")
    ET.SubElement(p_pr, qn("w", "jc"), {qn("w", "val"): "both" if alignment == "justify" else alignment})
    style_spacing = styles.get(style_id, {})
    ET.SubElement(
        p_pr,
        qn("w", "spacing"),
        {
            qn("w", "before"): str(_normalize_spacing(block.get("space_before"), int(style_spacing.get("space_before", 0))) * 20),
            qn("w", "after"): str(_normalize_spacing(block.get("space_after"), int(style_spacing.get("space_after", 0))) * 20),
            qn("w", "line"): str(int(round(_normalize_line_spacing(block.get("line_spacing"), float(style_spacing.get("line_spacing", DEFAULT_LINE_SPACING))) * 240))),
            qn("w", "lineRule"): "auto",
        },
    )
    numbering = _normalize_numbering(block.get("numbering"))
    if numbering is not None:
        num_id = list_id_to_num_id.get(numbering["list_id"] or "")
        if num_id:
            num_pr = ET.SubElement(p_pr, qn("w", "numPr"))
            ET.SubElement(num_pr, qn("w", "ilvl"), {qn("w", "val"): str(numbering["ilvl"])})
            ET.SubElement(num_pr, qn("w", "numId"), {qn("w", "val"): str(num_id)})

    runs = block.get("runs") or []
    style_descriptor = styles.get(style_id, {}).get("descriptor") or _default_descriptor()
    if not runs:
        paragraph.append(_run_node("", style_descriptor))
    for run_info in runs:
        paragraph.append(_run_node(run_info.get("text", ""), run_info.get("descriptor") or style_descriptor))
    return paragraph


def _table_node(block: dict, styles: dict[str, dict], list_id_to_num_id: dict[str, str]) -> ET.Element:
    table = ET.Element(qn("w", "tbl"))
    tbl_pr = ET.SubElement(table, qn("w", "tblPr"))
    ET.SubElement(tbl_pr, qn("w", "tblW"), {qn("w", "w"): "0", qn("w", "type"): "auto"})
    borders = ET.SubElement(tbl_pr, qn("w", "tblBorders"))
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        ET.SubElement(borders, qn("w", edge), {qn("w", "val"): "single", qn("w", "sz"): "4", qn("w", "space"): "0", qn("w", "color"): "C8B9A9"})

    for row in block.get("rows") or []:
        tr = ET.SubElement(table, qn("w", "tr"))
        for cell in row:
            tc = ET.SubElement(tr, qn("w", "tc"))
            tc_pr = ET.SubElement(tc, qn("w", "tcPr"))
            width = str(int(cell.get("width") or 2400))
            ET.SubElement(tc_pr, qn("w", "tcW"), {qn("w", "w"): width, qn("w", "type"): "dxa"})
            paragraphs = cell.get("paragraphs") or [{"type": "paragraph", "style_id": "Normal", "alignment": "align_left", "runs": []}]
            for paragraph in paragraphs:
                tc.append(_paragraph_node(paragraph, styles, list_id_to_num_id))
    return table


def _image_paragraph(rel_id: str, width_px: int, height_px: int, name: str) -> ET.Element:
    width_px = max(int(width_px or 320), 1)
    height_px = max(int(height_px or 180), 1)
    width_emu = int(width_px / 96 * 914400)
    height_emu = int(height_px / 96 * 914400)
    paragraph = ET.Element(qn("w", "p"))
    run = ET.SubElement(paragraph, qn("w", "r"))
    drawing = ET.SubElement(run, qn("w", "drawing"))
    inline = ET.SubElement(drawing, qn("wp", "inline"))
    ET.SubElement(inline, qn("wp", "extent"), {"cx": str(width_emu), "cy": str(height_emu)})
    ET.SubElement(inline, qn("wp", "docPr"), {"id": "1", "name": name or "image"})
    c_nv = ET.SubElement(inline, qn("wp", "cNvGraphicFramePr"))
    ET.SubElement(c_nv, qn("a", "graphicFrameLocks"), {"noChangeAspect": "1"})
    graphic = ET.SubElement(inline, qn("a", "graphic"))
    graphic_data = ET.SubElement(graphic, qn("a", "graphicData"), {"uri": "http://schemas.openxmlformats.org/drawingml/2006/picture"})
    pic = ET.SubElement(graphic_data, qn("pic", "pic"))
    nv_pic_pr = ET.SubElement(pic, qn("pic", "nvPicPr"))
    ET.SubElement(nv_pic_pr, qn("pic", "cNvPr"), {"id": "0", "name": name or "image"})
    ET.SubElement(nv_pic_pr, qn("pic", "cNvPicPr"))
    blip_fill = ET.SubElement(pic, qn("pic", "blipFill"))
    ET.SubElement(blip_fill, qn("a", "blip"), {qn("r", "embed"): rel_id})
    stretch = ET.SubElement(blip_fill, qn("a", "stretch"))
    ET.SubElement(stretch, qn("a", "fillRect"))
    sp_pr = ET.SubElement(pic, qn("pic", "spPr"))
    xfrm = ET.SubElement(sp_pr, qn("a", "xfrm"))
    ET.SubElement(xfrm, qn("a", "off"), {"x": "0", "y": "0"})
    ET.SubElement(xfrm, qn("a", "ext"), {"cx": str(width_emu), "cy": str(height_emu)})
    prst_geom = ET.SubElement(sp_pr, qn("a", "prstGeom"), {"prst": "rect"})
    ET.SubElement(prst_geom, qn("a", "avLst"))
    return paragraph


def _decode_meta_xml(value: object) -> bytes | None:
    if not value:
        return None
    try:
        return base64.b64decode(str(value))
    except Exception:
        return None


def _decode_meta_docx(value: object) -> bytes | None:
    payload = _decode_meta_xml(value)
    if payload is None:
        return None
    if payload[:2] != b"PK":
        return None
    return payload


def _meta_flag(value: object) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    return str(value).strip().lower() in {"1", "true", "yes", "on"}


_ALIGNMENT_TO_DOCX = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _ensure_docx_paragraph_style(doc, style_info: dict, style_name_by_id: dict[str, str]) -> str:
    style_id = style_info["id"]
    if style_id in style_name_by_id:
        return style_name_by_id[style_id]
    builtin_name = {
        "Normal": "Normal",
        "Heading1": "Heading 1",
        "Heading2": "Heading 2",
        "Heading3": "Heading 3",
    }.get(style_id)
    if builtin_name:
        style_name = builtin_name
        style = doc.styles[builtin_name]
    else:
        style_name = str(style_info.get("name") or style_id)
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    based_on = str(style_info.get("based_on") or "").strip()
    if based_on:
        base_name = style_name_by_id.get(based_on) or {
            "Normal": "Normal",
            "Heading1": "Heading 1",
            "Heading2": "Heading 2",
            "Heading3": "Heading 3",
        }.get(based_on)
        if base_name:
            try:
                style.base_style = doc.styles[base_name]
            except KeyError:
                pass
    descriptor = _normalize_descriptor(style_info.get("descriptor"))
    latin_font, _ = _split_font_family(descriptor[0])
    style.font.name = latin_font
    style.font.size = Pt(descriptor[1])
    style.font.bold = descriptor[2]
    style.font.italic = descriptor[3]
    style.font.underline = descriptor[4]
    _set_docx_rfonts(style.element.get_or_add_rPr(), descriptor[0])
    style.paragraph_format.alignment = _ALIGNMENT_TO_DOCX.get(str(style_info.get("alignment") or "left"), WD_ALIGN_PARAGRAPH.LEFT)
    style.paragraph_format.line_spacing = _normalize_line_spacing(style_info.get("line_spacing"), DEFAULT_LINE_SPACING)
    style.paragraph_format.space_before = Pt(_normalize_spacing(style_info.get("space_before"), 0))
    style.paragraph_format.space_after = Pt(_normalize_spacing(style_info.get("space_after"), 0))
    outline_level = style_info.get("outline_level")
    p_pr = style.element.get_or_add_pPr()
    outline_node = p_pr.find(docx_qn("w:outlineLvl"))
    if outline_level in {0, 1, 2}:
        if outline_node is None:
            outline_node = OxmlElement("w:outlineLvl")
            p_pr.append(outline_node)
        outline_node.set(docx_qn("w:val"), str(outline_level))
    elif outline_node is not None:
        p_pr.remove(outline_node)
    style_name_by_id[style_id] = style_name
    return style_name


def _set_docx_run_background(run, bg_color: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    shd = r_pr.find(docx_qn("w:shd"))
    if bg_color:
        if shd is None:
            shd = OxmlElement("w:shd")
            r_pr.append(shd)
        shd.set(docx_qn("w:val"), "clear")
        shd.set(docx_qn("w:color"), "auto")
        shd.set(docx_qn("w:fill"), bg_color.lstrip("#"))
    elif shd is not None:
        r_pr.remove(shd)


def _apply_docx_run_descriptor(run, descriptor: list | tuple | None) -> None:
    family, size, bold, italic, underline, bg_color = _normalize_descriptor(descriptor)
    latin_font, _ = _split_font_family(family)
    run.font.name = latin_font
    _set_docx_rfonts(run._element.get_or_add_rPr(), family)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    _set_docx_run_background(run, bg_color)


def _apply_docx_paragraph_numbering(paragraph, numbering: dict | None, list_id_to_num_id: dict[str, str]) -> None:
    numbering = _normalize_numbering(numbering)
    if numbering is None:
        return
    num_id = list_id_to_num_id.get(numbering["list_id"] or "")
    if not num_id:
        return

    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(docx_qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)

    ilvl = num_pr.find(docx_qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(docx_qn("w:val"), str(numbering["ilvl"]))

    num_id_node = num_pr.find(docx_qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(docx_qn("w:val"), str(num_id))


def _apply_docx_paragraph_format(paragraph, block: dict, style_info: dict | None) -> None:
    block_alignment = str(block.get("alignment") or "")
    if block_alignment == "align_center":
        alignment = "center"
    elif block_alignment == "align_right":
        alignment = "right"
    elif block_alignment == "align_justify":
        alignment = "justify"
    elif block_alignment == "align_left":
        alignment = "left"
    else:
        alignment = str((style_info or {}).get("alignment") or "left")
    paragraph.alignment = _ALIGNMENT_TO_DOCX.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    paragraph.paragraph_format.space_before = Pt(
        _normalize_spacing(block.get("space_before"), int((style_info or {}).get("space_before", 0)))
    )
    paragraph.paragraph_format.space_after = Pt(
        _normalize_spacing(block.get("space_after"), int((style_info or {}).get("space_after", 0)))
    )
    paragraph.paragraph_format.line_spacing = _normalize_line_spacing(
        block.get("line_spacing"), float((style_info or {}).get("line_spacing", DEFAULT_LINE_SPACING))
    )


def _write_docx_runs(paragraph, block: dict, style_info: dict | None) -> None:
    runs = block.get("runs") or []
    style_descriptor = (style_info or {}).get("descriptor") or _default_descriptor()
    if not runs:
        _apply_docx_run_descriptor(paragraph.add_run(""), style_descriptor)
        return
    for run_info in runs:
        run = paragraph.add_run(run_info.get("text", ""))
        _apply_docx_run_descriptor(run, run_info.get("descriptor") or style_descriptor)


def _write_docx_paragraph(container, doc, block: dict, styles: dict[str, dict], style_name_by_id: dict[str, str], list_id_to_num_id: dict[str, str]) -> None:
    paragraph = container.add_paragraph()
    style_id = _style_ref_for_block(block, styles)
    style_info = styles.get(style_id, styles.get("Normal"))
    if style_info is not None:
        paragraph.style = _ensure_docx_paragraph_style(doc, style_info, style_name_by_id)
    _apply_docx_paragraph_format(paragraph, block, style_info)
    _write_docx_runs(paragraph, block, style_info)
    _apply_docx_paragraph_numbering(paragraph, block.get("numbering"), list_id_to_num_id)


def _write_docx_table(doc, block: dict, styles: dict[str, dict], style_name_by_id: dict[str, str], list_id_to_num_id: dict[str, str]) -> None:
    rows = block.get("rows") or []
    column_count = max((len(row) for row in rows), default=1)
    table = doc.add_table(rows=max(len(rows), 1), cols=max(column_count, 1))
    for row_index, row in enumerate(rows):
        target_row = table.rows[row_index]
        for cell_index, cell_data in enumerate(row):
            cell = target_row.cells[cell_index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(docx_qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(docx_qn("w:type"), "dxa")
            tc_w.set(docx_qn("w:w"), str(int(cell_data.get("width") or 2400)))
            paragraphs = cell_data.get("paragraphs") or [{"type": "paragraph", "style_id": "Normal", "alignment": "align_left", "runs": []}]
            for paragraph_index, paragraph_block in enumerate(paragraphs):
                paragraph = cell.paragraphs[0] if paragraph_index == 0 else cell.add_paragraph()
                style_id = _style_ref_for_block(paragraph_block, styles)
                style_info = styles.get(style_id, styles.get("Normal"))
                if style_info is not None:
                    paragraph.style = _ensure_docx_paragraph_style(doc, style_info, style_name_by_id)
                _apply_docx_paragraph_format(paragraph, paragraph_block, style_info)
                _write_docx_runs(paragraph, paragraph_block, style_info)
                _apply_docx_paragraph_numbering(paragraph, paragraph_block.get("numbering"), list_id_to_num_id)


def _replace_docx_parts(source: bytes, replacements: dict[str, bytes]) -> bytes:
    memory = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(source), "r") as src, zipfile.ZipFile(memory, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        names = set(src.namelist())
        for item in src.infolist():
            if item.filename in replacements:
                dst.writestr(item, replacements[item.filename])
            else:
                dst.writestr(item, src.read(item.filename))
        for filename, payload in replacements.items():
            if filename not in names:
                dst.writestr(filename, payload)
    return memory.getvalue()


def _patch_docx_numbering(docx_bytes: bytes, numbering_xml: bytes | None) -> bytes:
    if numbering_xml is None:
        return docx_bytes
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        content_types = ET.fromstring(archive.read("[Content_Types].xml"))
        rels_root = ET.fromstring(archive.read("word/_rels/document.xml.rels"))

    content_type_tag = "{http://schemas.openxmlformats.org/package/2006/content-types}Override"
    has_override = any(node.attrib.get("PartName") == "/word/numbering.xml" for node in content_types.findall(content_type_tag))
    if not has_override:
        ET.SubElement(
            content_types,
            content_type_tag,
            {
                "PartName": "/word/numbering.xml",
                "ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            },
        )

    has_rel = any(
        rel.attrib.get("Type") == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering"
        for rel in rels_root.findall(qn("pr", "Relationship"))
    )
    if not has_rel:
        ET.SubElement(
            rels_root,
            qn("pr", "Relationship"),
            {
                "Id": "rIdNumbering",
                "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering",
                "Target": "numbering.xml",
            },
        )

    return _replace_docx_parts(
        docx_bytes,
        {
            "[Content_Types].xml": ET.tostring(content_types, encoding="utf-8", xml_declaration=True),
            "word/_rels/document.xml.rels": ET.tostring(rels_root, encoding="utf-8", xml_declaration=True),
            "word/numbering.xml": numbering_xml,
        },
    )


def document_to_docx_bytes(document: dict) -> bytes:
    blocks = document.get("blocks") or []
    styles = _style_map(_normalize_styles(document.get("styles")))
    meta = document.get("_docx_meta") if isinstance(document.get("_docx_meta"), dict) else {}
    preserved_docx = _decode_meta_docx(meta.get("source_docx_b64"))
    styles_dirty = _meta_flag(meta.get("styles_dirty"))
    numbering_dirty = _meta_flag(meta.get("numbering_dirty"))
    content_dirty = _meta_flag(meta.get("content_dirty"))
    page_dirty = _meta_flag(meta.get("page_dirty"))
    if preserved_docx is not None and not styles_dirty and not numbering_dirty and not content_dirty and not page_dirty:
        return preserved_docx
    list_id_to_num_id, numbering_xml = _build_numbering_payload(blocks)

    doc = WordDocument()
    style_name_by_id: dict[str, str] = {}
    for style in styles.values():
        _ensure_docx_paragraph_style(doc, style, style_name_by_id)

    section = doc.sections[0]
    page = document.get("page") or {}
    section.page_width = int(page.get("width_twips") or 11906)
    section.page_height = int(page.get("height_twips") or 16838)

    body = doc.element.body
    for child in list(body):
        if child.tag != docx_qn("w:sectPr"):
            body.remove(child)

    for block in blocks:
        block_type = block.get("type")
        if block_type == "image":
            paragraph = doc.add_paragraph()
            data_url = str(block.get("data_url") or "")
            if "," not in data_url:
                continue
            _, payload = data_url.split(",", 1)
            width_px = max(int(block.get("width_px") or 320), 1)
            height_px = max(int(block.get("height_px") or 180), 1)
            paragraph.add_run().add_picture(
                io.BytesIO(base64.b64decode(payload)),
                width=Inches(width_px / 96),
                height=Inches(height_px / 96),
            )
        elif block_type == "table":
            _write_docx_table(doc, block, styles, style_name_by_id, list_id_to_num_id)
        else:
            _write_docx_paragraph(doc, doc, block, styles, style_name_by_id, list_id_to_num_id)

    memory = io.BytesIO()
    doc.save(memory)
    return _patch_docx_numbering(memory.getvalue(), numbering_xml)


def _load_relationships(archive: zipfile.ZipFile) -> dict[str, str]:
    rel_map: dict[str, str] = {}
    try:
        rels_root = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
    except KeyError:
        return rel_map
    for rel in rels_root.findall(qn("pr", "Relationship")):
        rel_id = rel.attrib.get("Id")
        target = rel.attrib.get("Target")
        if rel_id and target:
            rel_map[rel_id] = f"word/{target}" if not target.startswith("word/") else target
    return rel_map


def _descriptor_from_properties(properties: ET.Element | None, fallback: list | None = None) -> list:
    descriptor = _normalize_descriptor(fallback)
    if properties is None:
        return descriptor
    r_fonts = properties.find(qn("w", "rFonts"))
    if r_fonts is not None:
        latin = r_fonts.attrib.get(qn("w", "ascii")) or r_fonts.attrib.get(qn("w", "hAnsi"))
        east_asia = r_fonts.attrib.get(qn("w", "eastAsia"))
        if latin and east_asia and latin != east_asia:
            descriptor[0] = f"{latin}, {east_asia}"
        else:
            descriptor[0] = east_asia or latin or descriptor[0]
    sz = properties.find(qn("w", "sz"))
    if sz is not None:
        descriptor[1] = max(int(sz.attrib.get(qn("w", "val"), str(descriptor[1] * 2))) // 2, 1)
    bold_node = properties.find(qn("w", "b"))
    italic_node = properties.find(qn("w", "i"))
    underline_node = properties.find(qn("w", "u"))
    if bold_node is not None:
        descriptor[2] = _toggle_property_enabled(bold_node)
    if italic_node is not None:
        descriptor[3] = _toggle_property_enabled(italic_node)
    if underline_node is not None:
        descriptor[4] = _toggle_property_enabled(underline_node, none_is_false=True)
    shd = properties.find(qn("w", "shd"))
    if shd is not None:
        fill = shd.attrib.get(qn("w", "fill"))
        if fill and fill.lower() != "auto":
            descriptor[5] = _normalize_color(f"#{fill}")
    highlight = properties.find(qn("w", "highlight"))
    if highlight is not None:
        val = highlight.attrib.get(qn("w", "val"))
        if val and val.lower() != "none":
            hex_color = _highlight_to_hex(val)
            if hex_color:
                descriptor[5] = hex_color
    return descriptor


def _parse_styles(archive: zipfile.ZipFile) -> list[dict]:
    try:
        root = ET.fromstring(archive.read("word/styles.xml"))
    except KeyError:
        return _builtin_styles()

    default_descriptor = _default_descriptor()
    defaults_root = root.find(qn("w", "docDefaults"))
    if defaults_root is not None:
        rpr_default = defaults_root.find(f"{qn('w', 'rPrDefault')}/{qn('w', 'rPr')}")
        if rpr_default is not None:
            default_descriptor = _descriptor_from_properties(rpr_default, default_descriptor)

    parsed: list[dict] = []
    for node in root.findall(qn("w", "style")):
        if node.attrib.get(qn("w", "type")) != "paragraph":
            continue
        raw_style_id = node.attrib.get(qn("w", "styleId")) or ""
        name_node = node.find(qn("w", "name"))
        raw_name = name_node.attrib.get(qn("w", "val")) if name_node is not None else ""
        style_id = _canonical_style_id(raw_style_id) or _canonical_style_id(raw_name)
        if style_id not in ALLOWED_STYLE_IDS:
            continue

        base_style = _style_map(_builtin_styles())[style_id]
        p_pr = node.find(qn("w", "pPr"))
        r_pr = node.find(qn("w", "rPr"))
        alignment = base_style.get("alignment", "left")
        line_spacing = base_style.get("line_spacing", DEFAULT_LINE_SPACING)
        space_before = base_style.get("space_before", 0)
        space_after = base_style.get("space_after", 0)
        style_numbering = None

        if p_pr is not None:
            jc = p_pr.find(qn("w", "jc"))
            if jc is not None:
                alignment = {
                    "both": "justify",
                    "distribute": "justify",
                    "justify": "justify",
                }.get(jc.attrib.get(qn("w", "val"), "left"), jc.attrib.get(qn("w", "val"), "left"))
            spacing = p_pr.find(qn("w", "spacing"))
            if spacing is not None:
                line_raw = spacing.attrib.get(qn("w", "line"))
                before_raw = spacing.attrib.get(qn("w", "before"))
                after_raw = spacing.attrib.get(qn("w", "after"))
                if line_raw:
                    line_spacing = max(int(line_raw) / 240, 1.0)
                if before_raw:
                    space_before = max(int(before_raw) // 20, 0)
                if after_raw:
                    space_after = max(int(after_raw) // 20, 0)
            num_pr = p_pr.find(qn("w", "numPr"))
            if num_pr is not None:
                num_id_node = num_pr.find(qn("w", "numId"))
                ilvl_node = num_pr.find(qn("w", "ilvl"))
                num_id = num_id_node.attrib.get(qn("w", "val")) if num_id_node is not None else None
                if num_id:
                    style_numbering = {
                        "num_id": str(num_id),
                        "ilvl": _normalize_ilvl(ilvl_node.attrib.get(qn("w", "val"), "0") if ilvl_node is not None else 0),
                    }

        parsed.append(
            _make_style(
                style_id,
                base_style["name"],
                _descriptor_from_properties(r_pr, default_descriptor),
                alignment,
                base_style.get("outline_level"),
                bool(base_style.get("is_default")),
                base_style.get("based_on"),
                line_spacing,
                space_before,
                space_after,
                style_numbering,
            )
        )
    return _normalize_styles({"paragraph": parsed})


def _parse_style_aliases(archive: zipfile.ZipFile) -> dict[str, str]:
    aliases: dict[str, str] = {}
    try:
        root = ET.fromstring(archive.read("word/styles.xml"))
    except KeyError:
        return aliases

    for node in root.findall(qn("w", "style")):
        if node.attrib.get(qn("w", "type")) != "paragraph":
            continue
        raw_style_id = node.attrib.get(qn("w", "styleId")) or ""
        if not raw_style_id:
            continue
        name_node = node.find(qn("w", "name"))
        raw_name = name_node.attrib.get(qn("w", "val")) if name_node is not None else ""
        canonical = _canonical_style_id(raw_style_id) or _canonical_style_id(raw_name)
        if canonical:
            aliases[raw_style_id] = canonical
    return aliases


def _merge_runs(runs: list[dict]) -> list[dict]:
    merged: list[dict] = []
    for run in runs:
        text = run.get("text", "")
        descriptor = _normalize_descriptor(run.get("descriptor"))
        if merged and merged[-1]["descriptor"] == descriptor:
            merged[-1]["text"] += text
        else:
            merged.append({"text": text, "descriptor": descriptor})
    return merged


def _numbering_from_num_id(num_id: str, ilvl: int, numbering_by_num_id: dict[str, dict[int, dict]]) -> dict:
    ilvl = _normalize_ilvl(ilvl)
    level_config = (numbering_by_num_id.get(num_id) or {}).get(
        ilvl,
        {"start": 1, "num_fmt": "decimal", "lvl_text": _default_lvl_text(ilvl)},
    )
    try:
        start = max(int(str(level_config.get("start", 1))), 1)
    except (TypeError, ValueError):
        start = 1
    return {
        "enabled": True,
        "list_id": f"num-{num_id}",
        "ilvl": ilvl,
        "num_fmt": _normalize_num_fmt(level_config.get("num_fmt")),
        "lvl_text": str(level_config.get("lvl_text") or _default_lvl_text(ilvl)),
        "start": start,
    }


def _resolve_style_numbering(style_id: str, styles_by_id: dict[str, dict]) -> dict | None:
    current = style_id
    visited: set[str] = set()
    for _ in range(24):
        if not current or current in visited:
            return None
        visited.add(current)
        style = styles_by_id.get(current)
        if not style:
            return None
        style_numbering = _normalize_style_numbering(style.get("numbering"))
        if style_numbering is not None:
            return style_numbering
        current = str(style.get("based_on") or "").strip()
    return None


def _parse_paragraph_node(
    paragraph: ET.Element,
    archive: zipfile.ZipFile,
    rel_map: dict[str, str],
    styles_by_id: dict[str, dict],
    style_aliases: dict[str, str],
    numbering_by_num_id: dict[str, dict[int, dict]],
) -> tuple[dict | None, list[dict]]:
    style_id = "Normal"
    clear_imported_format = False
    alignment = "align_left"
    p_pr = paragraph.find(qn("w", "pPr"))
    if p_pr is not None:
        p_style = p_pr.find(qn("w", "pStyle"))
        if p_style is not None:
            raw_style_id = p_style.attrib.get(qn("w", "val"), "Normal") or "Normal"
            canonical_style_id = _canonical_style_id(raw_style_id) or style_aliases.get(raw_style_id)
            if canonical_style_id:
                style_id = canonical_style_id
            else:
                style_id = "Normal"
                clear_imported_format = True
        jc = p_pr.find(qn("w", "jc"))
        if jc is not None:
            alignment = {
                "left": "align_left",
                "center": "align_center",
                "right": "align_right",
                "both": "align_justify",
                "distribute": "align_justify",
                "justify": "align_justify",
            }.get(jc.attrib.get(qn("w", "val"), "left"), "align_left")

    style_info = styles_by_id.get(style_id, styles_by_id.get("Normal"))
    line_spacing = DEFAULT_LINE_SPACING
    space_before = 0
    space_after = 0
    numbering = None
    if p_pr is None or p_pr.find(qn("w", "jc")) is None:
        alignment = {
            "left": "align_left",
            "center": "align_center",
            "right": "align_right",
            "justify": "align_justify",
        }.get((style_info or {}).get("alignment", "left"), "align_left")
    if p_pr is not None and not clear_imported_format:
        spacing = p_pr.find(qn("w", "spacing"))
        if spacing is not None:
            line_raw = spacing.attrib.get(qn("w", "line"))
            before_raw = spacing.attrib.get(qn("w", "before"))
            after_raw = spacing.attrib.get(qn("w", "after"))
            if line_raw:
                line_spacing = max(int(line_raw) / 240, 1.0)
            if before_raw:
                space_before = max(int(before_raw) // 20, 0)
            if after_raw:
                space_after = max(int(after_raw) // 20, 0)
        num_pr = p_pr.find(qn("w", "numPr"))
        if num_pr is not None:
            num_id_node = num_pr.find(qn("w", "numId"))
            ilvl_node = num_pr.find(qn("w", "ilvl"))
            num_id = num_id_node.attrib.get(qn("w", "val")) if num_id_node is not None else None
            if num_id:
                numbering = _numbering_from_num_id(
                    str(num_id),
                    ilvl_node.attrib.get(qn("w", "val"), "0") if ilvl_node is not None else 0,
                    numbering_by_num_id,
                )
    if numbering is None and not clear_imported_format:
        style_numbering = _resolve_style_numbering(style_id, styles_by_id)
        if style_numbering is not None:
            numbering = _numbering_from_num_id(
                style_numbering["num_id"],
                style_numbering["ilvl"],
                numbering_by_num_id,
            )

    runs: list[dict] = []
    images: list[dict] = []
    paragraph_descriptor = _default_descriptor() if clear_imported_format else ((style_info or {}).get("descriptor") or _default_descriptor())
    for run in paragraph.findall(qn("w", "r")):
        descriptor = _descriptor_from_properties(run.find(qn("w", "rPr")), paragraph_descriptor)
        for text_node in run.findall(qn("w", "t")):
            runs.append({"text": text_node.text or "", "descriptor": descriptor})
        for drawing in run.findall(qn("w", "drawing")):
            blip = drawing.find(f".//{qn('a', 'blip')}")
            extent = drawing.find(f".//{qn('wp', 'extent')}")
            if blip is None:
                continue
            rel_id = blip.attrib.get(qn("r", "embed"))
            media_path = rel_map.get(rel_id or "")
            if not media_path or media_path not in archive.namelist():
                continue
            binary = archive.read(media_path)
            ext = media_path.rsplit(".", 1)[-1].lower()
            mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp"}.get(ext, "image/png")
            width_px = 320
            height_px = 180
            if extent is not None:
                width_px = max(int(int(extent.attrib.get("cx", "3048000")) / 914400 * 96), 1)
                height_px = max(int(int(extent.attrib.get("cy", "1714500")) / 914400 * 96), 1)
            images.append({
                "type": "image",
                "name": media_path.split("/")[-1],
                "mime": mime,
                "data_url": f"data:{mime};base64,{base64.b64encode(binary).decode('ascii')}",
                "width_px": width_px,
                "height_px": height_px,
            })

    merged_runs = _merge_runs(runs)
    has_text = any(run["text"] for run in merged_runs)
    block = None
    if has_text or style_id != "Normal" or not images or numbering is not None:
        block = {
            "type": "paragraph",
            "style": _style_to_block_style(style_info),
            "style_id": style_id,
            "style_name": (style_info or {}).get("name", style_id),
            "alignment": alignment,
            "line_spacing": line_spacing,
            "space_before": space_before,
            "space_after": space_after,
            "runs": merged_runs,
        }
        if numbering is not None:
            block["numbering"] = numbering
        if clear_imported_format:
            block = _clear_imported_paragraph_format(block)
    return block, images


def _parse_table_node(
    table: ET.Element,
    archive: zipfile.ZipFile,
    rel_map: dict[str, str],
    styles_by_id: dict[str, dict],
    style_aliases: dict[str, str],
    numbering_by_num_id: dict[str, dict[int, dict]],
) -> dict:
    rows: list[list[dict]] = []
    for tr in table.findall(qn("w", "tr")):
        row_cells: list[dict] = []
        for tc in tr.findall(qn("w", "tc")):
            cell_paragraphs: list[dict] = []
            for paragraph in tc.findall(qn("w", "p")):
                block, _ = _parse_paragraph_node(paragraph, archive, rel_map, styles_by_id, style_aliases, numbering_by_num_id)
                if block is not None:
                    cell_paragraphs.append(block)
            if not cell_paragraphs:
                cell_paragraphs.append({"type": "paragraph", "style": "normal", "style_id": "Normal", "style_name": "Normal", "alignment": "align_left", "runs": []})
            width = 2400
            tc_pr = tc.find(qn("w", "tcPr"))
            if tc_pr is not None:
                tc_w = tc_pr.find(qn("w", "tcW"))
                if tc_w is not None:
                    width = int(tc_w.attrib.get(qn("w", "w"), "2400") or "2400")
            row_cells.append({"paragraphs": cell_paragraphs, "width": width})
        rows.append(row_cells)
    return {"type": "table", "rows": rows}


def docx_bytes_to_document(data: bytes) -> dict:
    blocks: list[dict] = []
    with zipfile.ZipFile(io.BytesIO(data), "r") as archive:
        meta: dict[str, str] = {}
        styles = _parse_styles(archive)
        styles_by_id = _style_map(styles)
        style_aliases = _parse_style_aliases(archive)
        rel_map = _load_relationships(archive)
        numbering_by_num_id = _parse_numbering(archive)
        root = ET.fromstring(archive.read("word/document.xml"))
        body = root.find(qn("w", "body"))
        if body is None:
            result = {"blocks": blocks, "styles": {"paragraph": styles}}
            if meta:
                result["_docx_meta"] = meta
            return result

        page_size = {"width_twips": 11906, "height_twips": 16838}
        sect_pr = body.find(qn("w", "sectPr"))
        if sect_pr is not None:
            pg_sz = sect_pr.find(qn("w", "pgSz"))
            if pg_sz is not None:
                page_size["width_twips"] = int(pg_sz.attrib.get(qn("w", "w"), page_size["width_twips"]))
                page_size["height_twips"] = int(pg_sz.attrib.get(qn("w", "h"), page_size["height_twips"]))

        for child in list(body):
            if child.tag == qn("w", "p"):
                block, images = _parse_paragraph_node(child, archive, rel_map, styles_by_id, style_aliases, numbering_by_num_id)
                if block is not None:
                    blocks.append(block)
                blocks.extend(images)
            elif child.tag == qn("w", "tbl"):
                blocks.append(_parse_table_node(child, archive, rel_map, styles_by_id, style_aliases, numbering_by_num_id))
    result = {"blocks": blocks, "styles": {"paragraph": styles}, "page": page_size}
    if meta:
        result["_docx_meta"] = meta
    return result
