import io
import unittest
import zipfile

from docx import Document

from docx_io import docx_bytes_to_document, document_to_docx_bytes


def paragraph_texts(document: dict) -> list[str]:
    return [
        "".join(run.get("text", "") for run in block.get("runs", []))
        for block in document["blocks"]
        if block.get("type") == "paragraph"
    ]


class LineBreakRoundTripTests(unittest.TestCase):
    def test_imports_word_soft_break(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("alpha\nbeta")
        buffer = io.BytesIO()
        doc.save(buffer)

        imported = docx_bytes_to_document(buffer.getvalue())

        self.assertEqual(paragraph_texts(imported), ["alpha", "beta"])

    def test_imports_word_carriage_return_element(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("alpha\nbeta")
        source = io.BytesIO()
        doc.save(source)

        rewritten = io.BytesIO()
        with zipfile.ZipFile(io.BytesIO(source.getvalue()), "r") as src, zipfile.ZipFile(rewritten, "w") as dst:
            for item in src.infolist():
                payload = src.read(item.filename)
                if item.filename == "word/document.xml":
                    payload = payload.replace(b"<w:br/>", b"<w:cr/>")
                dst.writestr(item, payload)

        imported = docx_bytes_to_document(rewritten.getvalue())

        self.assertEqual(paragraph_texts(imported), ["alpha", "beta"])

    def test_model_newline_survives_docx_round_trip(self):
        payload = {
            "blocks": [
                {
                    "type": "paragraph",
                    "style_id": "Normal",
                    "alignment": "align_left",
                    "runs": [
                        {
                            "text": "alpha\nbeta",
                            "descriptor": ["Times New Roman", 12, False, False, False],
                        }
                    ],
                }
            ]
        }

        exported = document_to_docx_bytes(payload)
        imported = docx_bytes_to_document(exported)

        self.assertEqual(paragraph_texts(imported), ["alpha", "beta"])


if __name__ == "__main__":
    unittest.main()
